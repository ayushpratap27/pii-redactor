import docx
import sys
import re
from typing import Dict, Tuple, List, Callable, Optional
from redactor import PIIRedactor

class DocxProcessor:
    """Parses and redacts DOCX documents preserving inline formatting and table layouts."""

    SPINNER_FRAMES = ['◐', '◓', '◑', '◒']

    def __init__(self, redactor: PIIRedactor, verbose: bool = True):
        self.redactor = redactor
        self.verbose = verbose

    def process_document(self, input_path: str, output_path: str, progress_callback: Optional[Callable[[int, int, str], None]] = None) -> Dict:
        doc = docx.Document(input_path)

        # Pre-calculate total work units for accurate process percentage
        total_body_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        total_table_paragraphs = sum(
            len(cell.paragraphs)
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        total_section_paragraphs = 0
        for section in doc.sections:
            for container in [section.header, section.first_page_header, section.even_page_header,
                              section.footer, section.first_page_footer, section.even_page_footer]:
                if container and not container.is_linked_to_previous:
                    total_section_paragraphs += len(container.paragraphs)

        total_steps = total_body_paragraphs + total_table_paragraphs + total_section_paragraphs
        if total_steps == 0:
            total_steps = 1

        current_step = 0
        total_paragraphs = 0
        processed_tables = 0
        total_redacted = 0
        category_counts: Dict[str, int] = {}

        def default_terminal_callback(step: int, total: int, phase: str):
            if not self.verbose:
                return
            pct = int((step / total) * 100) if total > 0 else 100
            frame = self.SPINNER_FRAMES[step % len(self.SPINNER_FRAMES)]
            sys.stdout.write(f"\r[{frame}] Redacting Document Process: {pct}%   ")
            sys.stdout.flush()
            if step >= total:
                sys.stdout.write("\n")
                sys.stdout.flush()

        callback = progress_callback or default_terminal_callback
        callback(0, total_steps, "Initializing Document AST...")

        # 1. Process Body Paragraphs
        for paragraph in doc.paragraphs:
            total_paragraphs += 1
            current_step += 1
            count, cats = self._redact_paragraph(paragraph)
            total_redacted += count
            for c, k in cats.items():
                category_counts[c] = category_counts.get(c, 0) + k
            callback(current_step, total_steps, f"Paragraph {total_paragraphs}/{total_body_paragraphs}")

        # 2. Process Tables
        for table in doc.tables:
            processed_tables += 1
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_step += 1
                        count, cats = self._redact_paragraph(paragraph)
                        total_redacted += count
                        for c, k in cats.items():
                            category_counts[c] = category_counts.get(c, 0) + k
                        callback(current_step, total_steps, f"Table {processed_tables}/{total_tables}")

        # 3. Process Headers and Footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and not header.is_linked_to_previous:
                    for paragraph in header.paragraphs:
                        current_step += 1
                        count, cats = self._redact_paragraph(paragraph)
                        total_redacted += count
                        for c, k in cats.items():
                            category_counts[c] = category_counts.get(c, 0) + k
                        callback(current_step, total_steps, "Headers")
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and not footer.is_linked_to_previous:
                    for paragraph in footer.paragraphs:
                        current_step += 1
                        count, cats = self._redact_paragraph(paragraph)
                        total_redacted += count
                        for c, k in cats.items():
                            category_counts[c] = category_counts.get(c, 0) + k
                        callback(current_step, total_steps, "Footers")

        callback(total_steps, total_steps, "Complete")
        doc.save(output_path)

        return {
            "total_paragraphs": total_paragraphs,
            "total_tables": total_tables,
            "total_redacted": total_redacted,
            "category_counts": category_counts,
            "mapping_count": len(self.redactor.entity_map)
        }

    def _is_structural_content(self, paragraph) -> bool:
        text = paragraph.text.strip()
        if not text:
            return True

        # Style check
        style_name = getattr(paragraph.style, 'name', '') if paragraph.style else ''
        style_upper = style_name.upper()
        if any(s in style_upper for s in ['TOC', 'TABLE OF CONTENTS', 'HEADING', 'TITLE', 'SUBTITLE', 'HEADER']):
            return True

        # Text-based TOC check (dot leaders or page numbers)
        if '....' in text or '…' in text or re.search(r'\.{3,}\s*\d+$', text):
            return True

        # Section Header check
        if text.isupper() and len(text.split()) <= 8 and any(kw in text for kw in ['SECTION', 'TABLE OF CONTENTS', 'CHAPTER', 'PART', 'ANNEXURE', 'EXHIBIT', 'INDEX']):
            return True

        # Glossary / Definition check
        if re.match(r'^(?:["“\'][^"“\'”]+["”\']|[A-Z0-9\s/._-]{2,30})\s*(?:means|includes|refers to|shall mean|shall include)\b', text, re.I):
            return True

        return False

    def _redact_paragraph(self, paragraph) -> Tuple[int, Dict[str, int]]:
        text = paragraph.text
        if not text or not text.strip() or self._is_structural_content(paragraph):
            return 0, {}

        spans = self.redactor.detect_entities(text)
        if not spans:
            return 0, {}

        cats: Dict[str, int] = {}
        for s in spans:
            t = s["type"]
            cats[t] = cats.get(t, 0) + 1

        self._apply_run_replacements(paragraph, spans)
        return len(spans), cats

    def _apply_run_replacements(self, paragraph, spans: List[Dict]):
        if not paragraph.runs or not spans:
            return

        # Snapshot run character boundaries BEFORE mutating text
        run_map = []
        curr = 0
        for run in paragraph.runs:
            rlen = len(run.text)
            run_map.append({
                "run": run,
                "start": curr,
                "end": curr + rlen,
                "text": run.text
            })
            curr += rlen

        sorted_spans = sorted(spans, key=lambda s: s["start"], reverse=True)
        for span in sorted_spans:
            replacement = self.redactor.get_replacement(span["text"], span["type"])
            start_offset = span["start"]
            end_offset = span["end"]

            for ritem in run_map:
                run = ritem["run"]
                r_start = ritem["start"]
                r_end = ritem["end"]

                if max(r_start, start_offset) < min(r_end, end_offset):
                    if start_offset >= r_start and end_offset <= r_end:
                        local_start = start_offset - r_start
                        local_end = end_offset - r_start
                        run.text = run.text[:local_start] + replacement + run.text[local_end:]
                    else:
                        local_start = max(0, start_offset - r_start)
                        local_end = min(r_end - r_start, end_offset - r_start)
                        if r_start <= start_offset < r_end:
                            run.text = run.text[:local_start] + replacement
                        elif r_start < end_offset <= r_end:
                            run.text = run.text[local_end:]
                        else:
                            run.text = ""
